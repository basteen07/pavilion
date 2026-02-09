from docx import Document
import os

samples = [
    r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Fitness Training\Fitness Equipment\Fitness Accessories\Fitness Accessories.docx',
    r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Individual Games\Badminton\Shuttlecock\Shuttlecock.docx',
    r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\More\Shoes\Cricket Shoes\Cricket Shoes.docx',
    r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Team Sports\Ball Games\Football\Football (Ball)\Football (Ball).docx'
]

output_file = 'sample_inspection.txt'

with open(output_file, 'w', encoding='utf-8') as f:
    for path in samples:
        f.write(f"\n--- FILE: {path} ---\n")
        try:
            doc = Document(path)
            f.write("Paragraphs (first 10):\n")
            for i, para in enumerate(doc.paragraphs[:10]):
                if para.text.strip():
                    f.write(f"{i}: {para.text.strip()}\n")
            
            f.write("\nTables:\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}:\n")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    f.write(str(row_data) + "\n")
        except Exception as e:
            f.write(f"Error reading: {e}\n")

print(f"Sample inspection complete. Results saved to {output_file}")
