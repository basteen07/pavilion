import pandas as pd
from docx import Document
import os
import re
import json

# Paths
found_products_file = 'found_products.json'
xlsx_path = r'c:\Users\dfuser\Desktop\pavilion-main\pavilion_advanced_template (20).xlsx'
output_xlsx = r'c:\Users\dfuser\Desktop\pavilion-main\pavilion_advanced_template_bulk.xlsx'

collection_map = {
    "Fitness Training": "FITNESS & TRAINING",
    "Individual Games": "INDIVIDUAL GAMES",
    "More": "MORE",
    "Team Sports": "TEAM SPORTS"
}

def slugify(text):
    if not text: return ""
    text = str(text).lower()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '-', text)
    return text.strip('-')

def clean_val(text):
    return str(text).strip()

def parse_docx(path):
    data = {
        'product_name': '',
        'brand': '',
        'sku': '',
        'options': [], # List of (name, value)
        'description': '',
        'short_description': ''
    }
    try:
        doc = Document(path)
        
        # Parse Tables for specific fields
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key = row.cells[0].text.strip().lower()
                    val = row.cells[1].text.strip()
                    
                    if 'product name' in key:
                        data['product_name'] = val
                    elif 'brand name' in key:
                        data['brand'] = val
                    elif 'code/model' in key or 'model no' in key:
                        data['sku'] = val.replace('FL:', '').replace('Code:', '').strip()
                    elif any(x in key for x in ['type', 'color', 'size', 'material', 'handle grip', 'style']):
                        data['options'].append((row.cells[0].text.strip(), val))

        # Paragraphs for Description
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paras:
            # Look for "Product Description:" header or just take first meaningful one
            desc_start = 0
            for i, p in enumerate(paras):
                if p.lower().startswith('product description'):
                    desc_start = i + 1
                    break
            
            data['short_description'] = paras[desc_start] if len(paras) > desc_start else paras[0]
            data['description'] = "\n".join(paras)
            
        # Fallback for SKU if still empty
        if not data['sku'] and data['product_name']:
            data['sku'] = "SKU-" + slugify(data['product_name'])[:20].upper()
            
    except Exception as e:
        print(f"Error parsing {path}: {e}")
    
    return data

def main():
    if not os.path.exists(found_products_file):
        print("found_products.json not found!")
        return

    with open(found_products_file, 'r', encoding='utf-8') as f:
        products_list = json.load(f)

    # Load template to get columns
    df_template = pd.read_excel(xlsx_path)
    final_rows = []

    print(f"Processing {len(products_list)} products...")

    for prod_info in products_list:
        docx_path = os.path.join(prod_info['root'], prod_info['docx'])
        extracted = parse_docx(docx_path)
        
        path_parts = prod_info['path_parts']
        
        # Collection
        collection = collection_map.get(path_parts[0], "MORE")
        
        # Category & Sub-Category
        category = "General"
        sub_category = ""
        
        if len(path_parts) >= 2:
            # E.g. Individual Games -> Badminton -> Racket
            # parts[1] is usually category
            category = path_parts[1]
            if len(path_parts) >= 3:
                sub_category = " > ".join(path_parts[2:])
        
        # Options
        opt1_name, opt1_val = "", ""
        opt2_name, opt2_val = "", ""
        if extracted['options']:
            opt1_name, opt1_val = extracted['options'][0]
            if len(extracted['options']) > 1:
                opt2_name, opt2_val = extracted['options'][1]

        # Build Row
        row = {
            'Product Handle (Optional)': slugify(extracted['product_name'] or prod_info['docx'].replace('.docx', '')),
            'Product Name *': extracted['product_name'] or prod_info['docx'].replace('.docx', ''),
            'SKU *': extracted['sku'] or "SKU-TBD",
            'Option1 Name': opt1_name,
            'Option1 Value': opt1_val,
            'Option2 Name': opt2_name,
            'Option2 Value': opt2_val,
            'Size': '',
            'Color': '',
            'MRP Price *': 0,
            'Dealer Price': 0,
            'Counter Price': 0,
            'Recommended Price': 0,
            'Shop Price': 0,
            'Collection': collection,
            'Category *': category,
            'Sub-Category': sub_category,
            'Tag': '',
            'Brand *': extracted['brand'] or 'Generic',
            'Description': extracted['description'],
            'Short Description': extracted['short_description'],
            'HSN Code': '',
            'Tax Class': 18,
            'Buy URL': '',
            'Unit/UoM': 1,
            'Images': ", ".join(prod_info['images'])
        }
        final_rows.append(row)

    print("Saving to Excel...")
    df_new_data = pd.DataFrame(final_rows)
    
    # Reorder columns to match template exactly
    for col in df_template.columns:
        if col not in df_new_data.columns:
            df_new_data[col] = ""
            
    df_final = pd.concat([df_template, df_new_data[df_template.columns]], ignore_index=True)
    
    df_final.to_excel(output_xlsx, index=False)
    print(f"Bulk update complete. {len(final_rows)} products added. File: {output_xlsx}")

if __name__ == "__main__":
    main()
