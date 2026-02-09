import pandas as pd
from docx import Document
import os
import sys

# Paths
docx_path = r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Individual Games\Table Tennis\Ply\Table Tennis Ply.docx'
xlsx_path = r'c:\Users\dfuser\Desktop\pavilion-main\pavilion_advanced_template (20).xlsx'
output_file = r'c:\Users\dfuser\Desktop\pavilion-main\inspection_results.txt'

with open(output_file, 'w', encoding='utf-8') as f:
    f.write("--- EXCEL HEADERS ---\n")
    try:
        df = pd.read_excel(xlsx_path)
        f.write(str(df.columns.tolist()) + "\n")
        f.write("\nFirst row of data:\n")
        f.write(str(df.head(1).to_dict(orient='records')) + "\n")
    except Exception as e:
        f.write(f"Error reading Excel: {e}\n")

    f.write("\n--- DOCX CONTENT ---\n")
    try:
        doc = Document(docx_path)
        f.write("Paragraphs:\n")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"{i}: {para.text.strip()}\n")
                
        f.write("\nTables:\n")
        for i, table in enumerate(doc.tables):
            f.write(f"Table {i}:\n")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                f.write(str(row_data) + "\n")
    except Exception as e:
        f.write(f"Error reading DOCX: {e}\n")

print(f"Inspection complete. Results saved to {output_file}")
