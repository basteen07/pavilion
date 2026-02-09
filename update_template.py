import pandas as pd
from docx import Document
import os
import re

# Resolve paths
docx_path = r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Individual Games\Table Tennis\Ply\Table Tennis Ply.docx'
images_dir = r'c:\Users\dfuser\Desktop\Pavilion Structure (Document & Images)\Individual Games\Table Tennis\Ply'
xlsx_path = r'c:\Users\dfuser\Desktop\pavilion-main\pavilion_advanced_template (20).xlsx'
output_xlsx = r'c:\Users\dfuser\Desktop\pavilion-main\pavilion_advanced_template_updated.xlsx'

def slugify(text):
    text = text.lower()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '-', text)
    return text.strip('-')

def parse_docx(path):
    doc = Document(path)
    data = {}
    
    # Extract from Tables
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        for row in table0.rows:
            key = row.cells[0].text.strip().lower()
            val = row.cells[1].text.strip()
            if 'product name' in key:
                data['product_name'] = val
            elif 'brand name' in key:
                data['brand'] = val
            elif 'code/model no' in key:
                data['sku'] = val.replace('FL:', '').strip()
            elif 'handle grip' in key:
                data['option1_name'] = 'Handle grip'
                data['option1_value'] = val

    # Extract Description / Short Description
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        data['short_description'] = paragraphs[1] if len(paragraphs) > 1 else paragraphs[0]
        data['description'] = "\n".join(paragraphs)
        
    return data

def main():
    print("Parsing DOCX...")
    extracted = parse_docx(docx_path)
    
    # Get Images
    print("Listing Images...")
    image_files = [f for f in os.listdir(images_dir) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    images_str = ", ".join(image_files)
    
    # Map to Excel columns
    new_row = {
        'Product Handle (Optional)': slugify(extracted.get('product_name', 'tt-ply')),
        'Product Name *': extracted.get('product_name', 'Table Tennis Ply'),
        'SKU *': extracted.get('sku', 'SKU-PENDING'),
        'Option1 Name': extracted.get('option1_name', ''),
        'Option1 Value': extracted.get('option1_value', ''),
        'Option2 Name': '',
        'Option2 Value': '',
        'Size': '',
        'Color': '',
        'MRP Price *': 0,
        'Dealer Price': 0,
        'Counter Price': 0,
        'Recommended Price': 0,
        'Shop Price': 0,
        'Collection': 'INDIVIDUAL GAMES',
        'Category *': 'Table Tennis',
        'Sub-Category': 'Ply',
        'Tag': '',
        'Brand *': extracted.get('brand', 'Generic'),
        'Description': extracted.get('description', ''),
        'Short Description': extracted.get('short_description', ''),
        'HSN Code': '',
        'Tax Class': 18,
        'Buy URL': '',
        'Unit/UoM': 1,
        'Images': images_str
    }
    
    print("Updating Excel Template...")
    try:
        df = pd.read_excel(xlsx_path)
        # Append new row
        df_new = pd.DataFrame([new_row])
        df_updated = pd.concat([df, df_new], ignore_index=True)
        
        # Save to new file to avoid corruption of original
        df_updated.to_excel(output_xlsx, index=False)
        print(f"Successfully updated Excel. Saved to: {output_xlsx}")
    except Exception as e:
        print(f"Error updating Excel: {e}")

if __name__ == "__main__":
    main()
